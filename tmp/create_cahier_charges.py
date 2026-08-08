from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path('/Users/sofiane/Documents/mikrotik')
OUT_DIR = ROOT / 'output' / 'documents'
TMP_DIR = ROOT / 'tmp' / 'cahier_charges'
OUT_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / 'Cahier_des_charges_MikroTik_Firewall_VPN.docx'
DIAGRAM_PATH = TMP_DIR / 'architecture.png'

TEAL = '0F766E'
TEAL_DARK = '134E4A'
TEAL_LIGHT = 'CCFBF1'
INK = '17212B'
MUTED = '52606D'
LIGHT = 'F4F7F8'
LINE = 'D8E1E5'
AMBER = 'F59E0B'
AMBER_LIGHT = 'FFF7ED'
RED = 'B91C1C'
RED_LIGHT = 'FEF2F2'
WHITE = 'FFFFFF'


def font(size, bold=False):
    candidates = [
        '/System/Library/Fonts/Supplemental/Arial Bold.ttf' if bold else '/System/Library/Fonts/Supplemental/Arial.ttf',
        '/System/Library/Fonts/Supplemental/Helvetica Bold.ttf' if bold else '/System/Library/Fonts/Supplemental/Helvetica.ttf',
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_architecture(path):
    canvas = Image.new('RGB', (1800, 720), '#F8FAFC')
    draw = ImageDraw.Draw(canvas)
    title_font = font(42, True)
    box_font = font(30, True)
    small_font = font(23, False)
    draw.text((70, 35), 'Architecture cible du laboratoire', fill='#134E4A', font=title_font)

    def box(x, y, w, h, fill, title, subtitle, outline='#D8E1E5'):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=outline, width=4)
        bbox = draw.textbbox((0, 0), title, font=box_font)
        draw.text((x + (w - (bbox[2] - bbox[0])) / 2, y + 28), title, fill='#17212B', font=box_font)
        lines = subtitle.split('\n')
        yy = y + 82
        for line in lines:
            bb = draw.textbbox((0, 0), line, font=small_font)
            draw.text((x + (w - (bb[2] - bb[0])) / 2, yy), line, fill='#52606D', font=small_font)
            yy += 32

    def arrow(x1, y1, x2, y2, label='', dashed=False):
        if dashed:
            steps = 15
            for i in range(0, steps, 2):
                xa = x1 + (x2 - x1) * i / steps
                ya = y1 + (y2 - y1) * i / steps
                xb = x1 + (x2 - x1) * (i + 1) / steps
                yb = y1 + (y2 - y1) * (i + 1) / steps
                draw.line((xa, ya, xb, yb), fill='#0F766E', width=6)
        else:
            draw.line((x1, y1, x2, y2), fill='#0F766E', width=7)
        draw.polygon([(x2, y2), (x2 - 20, y2 - 12), (x2 - 20, y2 + 12)], fill='#0F766E')
        if label:
            bb = draw.textbbox((0, 0), label, font=small_font)
            draw.rounded_rectangle(((x1+x2)/2 - (bb[2]-bb[0])/2 - 10, (y1+y2)/2 - 28,
                                    (x1+x2)/2 + (bb[2]-bb[0])/2 + 10, (y1+y2)/2 + 8),
                                   radius=8, fill='#FFFFFF')
            draw.text(((x1+x2)/2 - (bb[2]-bb[0])/2, (y1+y2)/2 - 24), label, fill='#134E4A', font=small_font)

    box(70, 235, 300, 165, '#E5E7EB', 'Internet / UTM', 'Shared Network\n192.168.64.0/24')
    box(525, 205, 360, 225, '#CCFBF1', 'R1-HQ', 'MikroTik CHR\nether1 = WAN\nether2 = LAN')
    box(1040, 235, 310, 165, '#FEF3C7', 'HQ-LAN', '10.10.10.0/24\nPasserelle 10.10.10.1')
    box(1490, 135, 250, 145, '#DBEAFE', 'Clients', 'Ubuntu\nWindows Server')
    box(1490, 385, 250, 145, '#EDE9FE', 'R2-Branch', 'Phase future\n10.20.20.0/24')
    box(530, 520, 350, 125, '#FCE7F3', 'Employes distants', 'VPN WireGuard individuel')

    arrow(370, 318, 525, 318, 'ether1 / WAN')
    arrow(885, 318, 1040, 318, 'ether2 / LAN')
    arrow(1350, 318, 1490, 215, 'DHCP')
    arrow(1350, 338, 1490, 455, 'VPN site a site', dashed=True)
    arrow(705, 520, 705, 430, 'VPN acces distant', dashed=True)
    canvas.save(path, quality=95)


draw_architecture(DIAGRAM_PATH)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.9)
section.right_margin = Cm(1.9)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size='8'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'start' if edge == 'left' else 'end' if edge == 'right' else edge
        el = borders.find(qn(f'w:{tag}'))
        if el is None:
            el = OxmlElement(f'w:{tag}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)


styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(10.2)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ('Title', 27, TEAL_DARK, 0, 8),
    ('Subtitle', 12, MUTED, 0, 12),
    ('Heading 1', 17, TEAL_DARK, 15, 7),
    ('Heading 2', 13.5, TEAL, 11, 5),
    ('Heading 3', 11.2, INK, 8, 3),
]:
    st = styles[name]
    st.font.name = 'Arial'
    st.font.size = Pt(size)
    st.font.bold = name != 'Subtitle'
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ('List Bullet', 'List Number'):
    st = styles[list_name]
    st.font.name = 'Arial'
    st.font.size = Pt(10.2)
    st.paragraph_format.left_indent = Cm(0.65)
    st.paragraph_format.first_line_indent = Cm(-0.35)
    st.paragraph_format.space_after = Pt(3.5)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = field_code
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char1, instr_text, fld_char2])


header = section.header
hp = header.paragraphs[0]
hp.text = 'PROJET MIKROTIK - FIREWALL ET VPN'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.runs[0].font.name = 'Arial'
hp.runs[0].font.size = Pt(8)
hp.runs[0].font.bold = True
hp.runs[0].font.color.rgb = RGBColor.from_string(TEAL)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Cahier des charges  |  ')
r.font.name = 'Arial'
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MUTED)
add_field(fp, 'PAGE')


def add_title(text, subtitle=None):
    p = doc.add_paragraph(style='Title')
    p.add_run(text)
    if subtitle:
        s = doc.add_paragraph(style='Subtitle')
        s.add_run(subtitle)
    return p


def add_p(text='', bold_prefix=None, align=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Cm(0.65 + 0.55 * level)
    p.add_run(text)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


def add_callout(title, text, fill=TEAL_LIGHT, border=TEAL):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.7)
    for cell in table.column_cells(0):
        cell.width = Cm(16.7)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 130, 190, 130, 190)
    set_table_borders(table, border, '10')
    p1 = table.cell(0, 0).paragraphs[0]
    p1.paragraph_format.space_after = Pt(1)
    rr = p1.add_run(title)
    rr.bold = True
    rr.font.name = 'Arial'
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = RGBColor.from_string(TEAL_DARK if fill != RED_LIGHT else RED)
    p2 = table.cell(1, 0).paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, (text, width) in enumerate(zip(headers, widths_cm)):
        cell = hdr.cells[i]
        cell.width = Cm(width)
        set_cell_shading(cell, TEAL)
        set_cell_margins(cell, 140, 150, 140, 150)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.name = 'Arial'
        r.font.size = Pt(9.1)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (text, width) in enumerate(zip(row, widths_cm)):
            cell = cells[i]
            cell.width = Cm(width)
            set_cell_margins(cell, 125, 150, 125, 150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ridx % 2:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            r.font.name = 'Arial'
            r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


# Cover page
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
r = p.add_run('CAHIER DES CHARGES')
r.font.name = 'Arial'
r.font.size = Pt(15)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(TEAL)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Mise en place d\'un firewall MikroTik et de VPN securises')
r.font.name = 'Arial'
r.font.size = Pt(29)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(TEAL_DARK)

p = doc.add_paragraph()
r = p.add_run('Laboratoire virtualise sous UTM sur Apple Silicon')
r.font.name = 'Arial'
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(MUTED)

doc.add_paragraph().add_run('')
add_callout(
    'Finalite du projet',
    'Construire une plateforme de securite reseau permettant de filtrer le trafic du siege, de relier deux routeurs et de donner aux employes un acces distant chiffre et controle.',
)

doc.add_picture(str(DIAGRAM_PATH), width=Cm(16.7))
last = doc.paragraphs[-1]
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = add_table(
    ['Champ', 'Valeur'],
    [
        ('Projet', 'Firewall et VPN MikroTik CHR'),
        ('Routeur principal', 'R1-HQ'),
        ('Plateforme', 'UTM / QEMU / Apple Silicon'),
        ('Version du document', '1.0'),
        ('Date', '17 juillet 2026'),
        ('Statut', 'Etapes 1 et 2 terminees - Etape 3 en cours'),
    ],
    [4.2, 12.5],
)

doc.add_page_break()

add_title('Synthese executive', 'Vision, perimetre et etat actuel du projet')
add_callout(
    'Decision d\'architecture',
    'R1-HQ est place entre un cote WAN non fiable et un reseau interne HQ-LAN. Tout acces Internet des clients doit traverser le MikroTik. Les acces distants et la future liaison avec R2-Branch utiliseront WireGuard.',
)

doc.add_heading('1. Contexte et besoin', level=1)
add_p('L\'entreprise souhaite disposer d\'un firewall capable de controler les connexions entre Internet, le reseau du siege et un futur site distant. Le point le plus important concerne les employes qui doivent se connecter depuis l\'exterieur sans exposer directement les services internes.')
add_p('Le projet est d\'abord construit dans un laboratoire UTM afin de tester chaque regle avant une presentation ou un deploiement reel. Cette methode reduit le risque de coupure, de perte d\'acces administratif ou de regle trop permissive.')

doc.add_heading('2. Objectifs generaux', level=1)
for item in [
    'Creer un routeur/firewall principal R1-HQ sous MikroTik CHR.',
    'Separer clairement le WAN, le LAN du siege et les futurs reseaux distants.',
    'Fournir DHCP, DNS et NAT aux machines du reseau interne.',
    'Mettre en oeuvre une politique de firewall restrictive et testable.',
    'Deployer un VPN WireGuard pour les employes distants.',
    'Preparer un VPN site a site entre R1-HQ et un futur R2-Branch.',
    'Tester le comportement avec Ubuntu, Windows Server, pfSense et des outils de diagnostic.',
    'Documenter les configurations, sauvegardes, tests et preuves de fonctionnement.',
]:
    add_bullet(item)

doc.add_heading('3. Perimetre', level=1)
doc.add_heading('3.1 Inclus dans le projet', level=2)
for item in [
    'Virtualisation UTM sur Mac Apple Silicon.',
    'MikroTik CHR ARM64 comme routeur et firewall.',
    'Reseaux virtuels WAN, HQ-LAN et futur BRANCH-LAN.',
    'Services DHCP, DNS, NAT, firewall, journalisation et VPN.',
    'Machines clientes de test et scenarios d\'acceptation.',
    'Backups RouterOS et documentation technique.',
]:
    add_bullet(item)
doc.add_heading('3.2 Hors perimetre actuel', level=2)
for item in [
    'Deploiement direct sur le reseau de production de l\'entreprise.',
    'Achat de routeurs physiques, licences commerciales ou certificats publics.',
    'Haute disponibilite avec deux firewalls actifs/passifs.',
    'Authentification multifacteur native WireGuard; une solution RADIUS/IdP complementaire devra etre etudiee si le MFA est obligatoire.',
]:
    add_bullet(item)

add_title('Architecture et exigences', 'Description technique de la cible')

doc.add_heading('4. Architecture cible', level=1)
doc.add_picture(str(DIAGRAM_PATH), width=Cm(16.7))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_p('Le trafic suit la chaine suivante : client interne -> ether2 -> R1-HQ -> controles firewall/NAT -> ether1 -> Internet. Les connexions entrantes du WAN sont refusees par defaut, sauf les services explicitement autorises, notamment le port du VPN.')

doc.add_heading('5. Plan d\'adressage', level=1)
add_table(
    ['Zone', 'Reseau / adresse', 'Passerelle', 'Usage'],
    [
        ('WAN UTM', '192.168.64.0/24', '192.168.64.1', 'Acces externe du laboratoire'),
        ('R1-HQ WAN', '192.168.64.2', '192.168.64.1', 'WebFig et sortie Internet'),
        ('HQ-LAN', '10.10.10.0/24', '10.10.10.1', 'Reseau du siege'),
        ('Pool DHCP HQ', '10.10.10.100-200', '10.10.10.1', 'Postes clients dynamiques'),
        ('Ubuntu test', '10.10.10.200/24', '10.10.10.1', 'Client obtenu par DHCP'),
        ('BRANCH-LAN', '10.20.20.0/24', '10.20.20.1', 'Phase future R2-Branch'),
        ('VPN employes', '10.99.0.0/24', 'R1-HQ', 'Adresses WireGuard proposees'),
    ],
    [3.2, 4.1, 3.6, 5.8],
)

doc.add_heading('6. Exigences fonctionnelles', level=1)
reqs = [
    ('F-01', 'R1-HQ doit router le trafic entre HQ-LAN et le WAN.'),
    ('F-02', 'Les clients HQ-LAN doivent recevoir automatiquement une adresse, une passerelle et un DNS.'),
    ('F-03', 'Le NAT doit traduire uniquement le trafic sortant autorise de 10.10.10.0/24.'),
    ('F-04', 'Le firewall doit refuser les connexions entrantes non sollicitees depuis le WAN.'),
    ('F-05', 'L\'administration du routeur doit etre limitee aux reseaux autorises.'),
    ('F-06', 'Chaque employe distant doit disposer de son propre pair WireGuard et de sa propre cle.'),
    ('F-07', 'R1-HQ et R2-Branch doivent echanger uniquement les sous-reseaux approuves via un tunnel site a site.'),
    ('F-08', 'Les journaux doivent permettre d\'identifier les connexions refusees importantes sans saturer le routeur.'),
    ('F-09', 'Une sauvegarde doit etre creee avant et apres chaque changement majeur.'),
]
add_table(['ID', 'Exigence'], reqs, [2.0, 14.7])

doc.add_heading('7. Exigences de securite', level=1)
for item in [
    'Politique par defaut restrictive : autoriser ce qui est necessaire, refuser le reste.',
    'Regles basees sur les etats de connexion : established/related autorises, invalid refuses.',
    'Aucun service d\'administration expose inutilement sur le WAN.',
    'Telnet, FTP, API, API-SSL et Reverse Proxy desactives tant qu\'ils ne sont pas justifies.',
    'WebFig, SSH et WinBox conserves pour l\'administration, puis limites au LAN/VPN.',
    'Cles WireGuard individuelles, revocables et non partagees entre employes.',
    'Principe du moindre privilege pour l\'acces VPN aux serveurs internes.',
    'Backups conserves hors du routeur et mots de passe non inclus dans les rapports.',
]:
    add_bullet(item)

doc.add_page_break()
add_title('Dossier de realisation', 'Travaux executes depuis la creation du laboratoire')

doc.add_heading('8. Etape 0 - Creation de la VM MikroTik CHR', level=1)
add_p('Une nouvelle VM QEMU nommee MikroTik CHR Automated a ete creee dans /Applications/UTM.app sans modifier ni supprimer les autres machines UTM existantes. Le disque ARM64 officiel chr-7.22.1-arm64.qcow2 a ete verifie avant import afin d\'eviter la creation accidentelle d\'un disque vide.')
add_table(
    ['Parametre', 'Configuration'],
    [
        ('Architecture', 'aarch64 / ARM64'),
        ('Machine QEMU', 'virt ARM virtual machine'),
        ('Memoire', '2048 MiB'),
        ('CPU', '2 coeurs'),
        ('Acceleration', 'Apple Hypervisor active si supporte'),
        ('Demarrage', 'UEFI'),
        ('Affichage / son', 'Aucun affichage graphique, aucun son'),
        ('Console', 'Une console serie PTTY'),
        ('WAN initial', 'Shared Network, virtio-net-pci'),
        ('Disque', 'Image MikroTik ARM64 importee comme VirtIO non amovible'),
    ],
    [5.1, 11.6],
)

doc.add_heading('9. Etape 1 - Preparation et securisation de R1-HQ', level=1)
for item in [
    'Connexion WebFig a http://192.168.64.2.',
    'Renommage du routeur MikroTik vers R1-HQ : R1 pour routeur principal, HQ pour Headquarters.',
    'Creation du backup R1-HQ-baseline.backup avant les changements.',
    'Desactivation de Telnet, FTP, API, API-SSL et Reverse Proxy.',
    'Conservation de WebFig, SSH et WinBox pour l\'administration.',
    'Creation du backup R1-HQ-step1-secured.backup apres durcissement.',
]:
    add_bullet(item)
add_callout('Pourquoi le firewall a ete reporte', 'A ce moment, le routeur ne possedait que ether1. Il fallait d\'abord creer ether2 pour differencier sans ambiguite le WAN non fiable du LAN interne et eviter de bloquer WebFig.', AMBER_LIGHT, AMBER)

doc.add_heading('10. Etape 2 - Creation du LAN du siege', level=1)
for item in [
    'Arret propre de R1-HQ avant modification du materiel virtuel.',
    'Creation du Host Network UTM HQ-LAN.',
    'Ajout d\'une seconde carte virtio-net-pci en mode macOS Host Only.',
    'Association de ether2 au reseau HQ-LAN.',
    'Attribution de 10.10.10.1/24 a ether2 avec le commentaire HQ LAN Gateway.',
    'Creation du pool DHCP 10.10.10.100-10.10.10.200.',
    'Configuration du reseau DHCP 10.10.10.0/24 et de la passerelle 10.10.10.1.',
]:
    add_bullet(item)
add_callout('Incident Safe Mode', 'Le message setup failed to enter safe mode: safe mode already held by somebody (8) provenait d\'un Safe Mode deja actif dans WebFig. La solution etait de liberer ce mode avant de relancer DHCP Setup.', RED_LIGHT, RED)

doc.add_heading('11. Etape 3 - NAT, DNS et client Ubuntu', level=1)
add_p('Cette etape valide que R1-HQ est reellement place sur le chemin du trafic et qu\'un client du LAN peut utiliser ses services.')
add_table(
    ['Element', 'Valeur / resultat', 'Statut'],
    [
        ('Regle NAT', 'srcnat masquerade, source 10.10.10.0/24, sortie ether1', 'Valide'),
        ('Compteurs observes', '115.7 KiB et 1 532 paquets au moment du controle', 'Preuve de passage'),
        ('DNS MikroTik', 'Allow Remote Requests active', 'Valide'),
        ('DNS distribue', '10.10.10.1 via DHCP', 'Valide'),
        ('VM cliente', 'HQ-Test-Client, Ubuntu Server ARM64', 'Creee'),
        ('Disque client', '20 Go VirtIO, nouvelle VM uniquement', 'Pret'),
        ('Adresse Ubuntu', '10.10.10.200/24 par DHCP', 'Valide'),
        ('Miroir Ubuntu', 'ma.ports.ubuntu.com, test reussi', 'Valide'),
        ('Telechargement test', '508 kB, environ 135 kB/s', 'Valide'),
        ('Installation Ubuntu', 'Ecran Guided storage atteint, ecriture disque en attente d\'approbation', 'En cours'),
    ],
    [3.8, 9.8, 3.1],
)
doc.add_page_break()
add_title('Politique firewall et VPN', 'Regles cibles a implementer et a tester')

doc.add_heading('12. Politique de filtrage cible', level=1)
doc.add_heading('12.1 Chaine input - protection du routeur', level=2)
input_rules = [
    ('1', 'Accept', 'established, related', 'Maintenir les connexions legitimes'),
    ('2', 'Drop', 'invalid', 'Refuser les paquets incoherents'),
    ('3', 'Accept', 'ICMP limite', 'Diagnostic sans abus'),
    ('4', 'Accept', 'HQ-LAN -> administration', 'WebFig, SSH et WinBox depuis le LAN'),
    ('5', 'Accept', 'VPN WireGuard depuis WAN', 'Autoriser uniquement le port du tunnel'),
    ('6', 'Drop', 'tout le reste depuis WAN', 'Politique restrictive'),
]
add_table(['Ordre', 'Action', 'Condition', 'But'], input_rules, [1.6, 2.3, 6.1, 6.7])

doc.add_heading('12.2 Chaine forward - protection des reseaux', level=2)
forward_rules = [
    ('1', 'Accept', 'established, related', 'Reponses aux connexions autorisees'),
    ('2', 'Drop', 'invalid', 'Trafic invalide'),
    ('3', 'Accept', 'HQ-LAN -> WAN', 'Internet pour le siege'),
    ('4', 'Accept limite', 'VPN employes -> serveurs autorises', 'Moindre privilege'),
    ('5', 'Accept', 'HQ-LAN <-> BRANCH-LAN', 'Uniquement via WireGuard site a site'),
    ('6', 'Drop', 'WAN -> LAN non sollicite', 'Bloquer les connexions entrantes'),
    ('7', 'Drop', 'tout le reste', 'Fin de politique explicite'),
]
add_table(['Ordre', 'Action', 'Condition', 'But'], forward_rules, [1.6, 2.3, 6.1, 6.7])

doc.add_heading('13. VPN d\'acces distant pour les employes', level=1)
for item in [
    'Technologie recommandee : WireGuard pour sa simplicite, ses performances et sa cryptographie moderne.',
    'Une cle privee et une adresse VPN uniques par employe; aucune cle partagee.',
    'Le pare-feu autorise chaque groupe d\'employes uniquement vers les ressources necessaires.',
    'Un pair peut etre desactive immediatement en cas de depart, perte de poste ou suspicion.',
    'Les configurations clientes sont transmises par un canal securise et ne sont jamais publiees dans le rapport.',
    'Le MFA n\'est pas fourni nativement par WireGuard; si l\'entreprise l\'exige, ajouter une passerelle d\'identite/RADIUS ou une solution VPN avec authentification forte.',
]:
    add_bullet(item)

doc.add_heading('14. VPN site a site R1-HQ - R2-Branch', level=1)
for item in [
    'Creer R2-Branch dans une VM separee sans reutiliser les disques de R1-HQ.',
    'Configurer BRANCH-LAN en 10.20.20.0/24 avec la passerelle 10.20.20.1.',
    'Creer un tunnel WireGuard dedie entre les deux routeurs.',
    'Limiter Allowed Address aux reseaux 10.10.10.0/24 et 10.20.20.0/24.',
    'Ajouter les routes et regles forward necessaires, sans NAT entre les deux LAN.',
    'Tester ping, DNS, acces aux services autorises et blocage des services interdits.',
]:
    add_number(item)

doc.add_page_break()
add_title('Tests, livrables et planning', 'Conditions de validation du projet')

doc.add_heading('15. Plan de tests', level=1)
tests = [
    ('T-01', 'DHCP', 'Le client recoit une adresse 10.10.10.100-200, gateway 10.10.10.1 et DNS 10.10.10.1.'),
    ('T-02', 'Passerelle', 'Le client joint 10.10.10.1.'),
    ('T-03', 'Internet', 'Le client joint une adresse Internet et le compteur NAT augmente.'),
    ('T-04', 'DNS', 'Le client resout un nom public via R1-HQ.'),
    ('T-05', 'WAN vers LAN', 'Une connexion non sollicitee depuis le WAN est bloquee.'),
    ('T-06', 'Administration', 'WebFig/SSH/WinBox sont accessibles depuis le LAN ou VPN autorise et bloques ailleurs.'),
    ('T-07', 'VPN employe', 'Un pair valide rejoint uniquement les ressources autorisees.'),
    ('T-08', 'Revocation', 'Un pair desactive ne peut plus se connecter.'),
    ('T-09', 'Site a site', 'HQ-LAN et BRANCH-LAN communiquent uniquement selon la matrice d\'acces.'),
    ('T-10', 'Journalisation', 'Les refus importants sont visibles sans flood excessif des logs.'),
    ('T-11', 'Restauration', 'Un backup RouterOS peut restaurer l\'etape correspondante.'),
]
add_table(['ID', 'Test', 'Critere d\'acceptation'], tests, [1.8, 3.2, 11.7])

doc.add_heading('16. Livrables', level=1)
for item in [
    'VM MikroTik CHR Automated configuree comme R1-HQ.',
    'VM cliente HQ-Test-Client sous Ubuntu Server.',
    'Backups R1-HQ-baseline.backup, R1-HQ-step1-secured.backup et R1-HQ-step2-LAN.backup.',
    'Backups supplementaires apres NAT, firewall et VPN.',
    'Fichiers Etape 1, Etape 2, cahier des charges et futur rapport de stage.',
    'Plan d\'adressage, matrice firewall, inventaire des peers VPN sans cles privees.',
    'Captures et resultats de tests prouvant le fonctionnement et le blocage attendu.',
]:
    add_bullet(item)

doc.add_heading('17. Planning de realisation sur 8 jours', level=1)
planning = [
    ('Jour 1', 'VM CHR, identification R1-HQ, backups, services inutiles desactives', 'Termine'),
    ('Jour 2', 'HQ-LAN, ether2, adressage, DHCP', 'Termine'),
    ('Jour 3', 'NAT, DNS local, VM Ubuntu, validation Internet', 'En cours'),
    ('Jour 4', 'Installation Ubuntu, firewall input et tests de non-regression', 'A faire'),
    ('Jour 5', 'Firewall forward, logs, tests autorise/refuse', 'A faire'),
    ('Jour 6', 'WireGuard employes, pairs individuels, revocation', 'A faire'),
    ('Jour 7', 'R2-Branch, VPN site a site, Windows Server/pfSense', 'A faire'),
    ('Jour 8', 'Tests finaux, backups, captures, presentation et rapport', 'A faire'),
]
add_table(['Periode', 'Travaux', 'Etat'], planning, [2.2, 11.6, 2.9])

doc.add_heading('18. Risques et mesures de reduction', level=1)
risks = [
    ('Blocage de WebFig', 'Eleve', 'Safe Mode, acces serie, backup avant les regles drop.'),
    ('Confusion WAN/LAN', 'Eleve', 'Noms explicites ether1-WAN et ether2-LAN, interface lists.'),
    ('Regle trop permissive', 'Eleve', 'Tests par cas, principe du moindre privilege, drop final.'),
    ('Perte de cle VPN', 'Eleve', 'Peer individuel, revocation immediate, inventaire controle.'),
    ('DNS indisponible', 'Moyen', 'R1-HQ comme resolver LAN, test DNS separe du ping IP.'),
    ('Echec VM ou disque', 'Moyen', 'Images verifiees, backups UTM/RouterOS, aucun disque vide cree.'),
    ('Journalisation excessive', 'Moyen', 'Rate limit et journalisation ciblee.'),
]
add_table(['Risque', 'Impact', 'Mesure'], risks, [4.1, 2.5, 10.1])

doc.add_heading('19. Criteres de recette finale', level=1)
for item in [
    'R1-HQ demarre de maniere fiable et conserve sa configuration.',
    'Un client HQ-LAN obtient automatiquement sa configuration reseau.',
    'Internet et DNS fonctionnent uniquement en passant par R1-HQ.',
    'Les connexions WAN non autorisees sont bloquees.',
    'L\'administration est limitee aux sources autorisees.',
    'Un employe distant autorise atteint uniquement les services prevus.',
    'Un pair VPN revoque est effectivement bloque.',
    'La liaison R1-HQ/R2-Branch respecte la matrice des flux.',
    'Les backups, configurations et preuves de test sont disponibles.',
]:
    add_bullet(item)

add_callout('Etat a la date du document', 'Les fondations reseau sont operationnelles et les tests DHCP/DNS/NAT ont reussi. L\'installation Ubuntu est arrivee a l\'ecran de partitionnement du disque de la nouvelle VM; elle reste a confirmer avant de poursuivre vers le firewall et WireGuard.', AMBER_LIGHT, AMBER)

doc.core_properties.title = 'Cahier des charges - Firewall et VPN MikroTik'
doc.core_properties.subject = 'Laboratoire UTM, MikroTik CHR, firewall et WireGuard'
doc.core_properties.author = 'Projet MikroTik'
doc.core_properties.keywords = 'MikroTik, CHR, UTM, firewall, WireGuard, VPN, DHCP, NAT'

doc.save(DOCX_PATH)
print(DOCX_PATH)
