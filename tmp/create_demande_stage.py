from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "/Users/sofiane/Documents/mikrotik/output/documents/Demande_de_stage_El_Houssaimi_Sofiane.docx"


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


doc = Document()
section = doc.sections[0]

# Named layout override for a Moroccan administrative letter: A4 instead of Letter.
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.10

# Quiet header, matching the standard_business_brief preset.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("DEMANDE DE STAGE")
set_font(hr, size=8.5, bold=True, color="667085")

# Sender + date block.
top = doc.add_table(rows=1, cols=2)
top.alignment = WD_TABLE_ALIGNMENT.CENTER
top.autofit = False
top.columns[0].width = Cm(8.1)
top.columns[1].width = Cm(8.1)
remove_table_borders(top)

left = top.cell(0, 0)
right = top.cell(0, 1)
left.width = Cm(8.1)
right.width = Cm(8.1)
left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
set_cell_margins(left, 0, 0, 0, 80)
set_cell_margins(right, 0, 80, 0, 0)

p = left.paragraphs[0]
p.paragraph_format.space_after = Pt(2)
r = p.add_run("El Houssaimi Sofiane")
set_font(r, size=11.5, bold=True, color="17365D")
for label, value in (
    ("CIN", "[À compléter]"),
    ("Tél.", "[À compléter]"),
    ("E-mail", "[À compléter]"),
):
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{label} : ")
    set_font(r, size=10.5, bold=True, color="344054")
    r = p.add_run(value)
    set_font(r, size=10.5, color="344054")

p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Agadir, le 06 août 2026")
set_font(r, size=10.5, color="344054")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Recipient block.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(1)
r = p.add_run("À l’attention de Monsieur le Responsable informatique")
set_font(r, size=10.5, bold=True, color="344054")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Sté Abhaje Frères")
set_font(r, size=10.5, bold=True, color="17365D")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Agadir, Maroc")
set_font(r, size=10.5, color="344054")

# Formal application heading.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("DEMANDE DE STAGE")
set_font(r, size=18, bold=True, color="17365D")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(15)
r = p.add_run("Objet : Formalisation d’un stage en Génie Informatique")
set_font(r, size=11, bold=True, color="44546A")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(9)
r = p.add_run("Monsieur,")
set_font(r, size=11)

paragraphs = [
    (
        "Étudiant en Génie Informatique à UNIVERSIAPOLIS Agadir, je me permets de vous "
        "adresser la présente demande afin de formaliser mon stage pratique d’une durée "
        "d’un mois et demi au sein de la Sté Abhaje Frères, pour la période allant du "
        "[date de début] au [date de fin]."
    ),
    (
        "Ce stage s’inscrit dans le cadre de ma formation et porte principalement sur la "
        "conception et la mise en place d’une solution de sécurité réseau : configuration "
        "d’un pare-feu MikroTik, filtrage des communications, sécurisation des accès "
        "distants par VPN et interconnexion sécurisée de plusieurs réseaux."
    ),
    (
        "Cette expérience me permet de mettre en pratique mes connaissances en réseaux "
        "informatiques et en cybersécurité, tout en développant mon autonomie, ma rigueur "
        "et ma capacité à répondre aux besoins techniques réels de l’entreprise."
    ),
    (
        "Je vous remercie de l’attention portée à ma demande et reste à votre disposition "
        "pour toute information ou formalité complémentaire."
    ),
]

for text in paragraphs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_font(r, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(0.7)
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(15)
r = p.add_run(
    "Dans l’attente de votre accord, je vous prie d’agréer, Monsieur, l’expression "
    "de ma considération distinguée."
)
set_font(r, size=11)

# Signature block with restrained fill.
sig = doc.add_table(rows=1, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
sig.autofit = False
sig.columns[0].width = Cm(9.5)
sig.columns[1].width = Cm(6.7)
remove_table_borders(sig)
sig.cell(0, 0).width = Cm(9.5)
sig.cell(0, 1).width = Cm(6.7)
set_cell_margins(sig.cell(0, 0), 0, 0, 0, 80)
set_cell_margins(sig.cell(0, 1), 80, 180, 80, 180)
shade_cell(sig.cell(0, 1), "F2F4F7")

p = sig.cell(0, 1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("Signature du stagiaire")
set_font(r, size=9.5, bold=True, color="44546A")
p = sig.cell(0, 1).add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("El Houssaimi Sofiane")
set_font(r, size=10.5, bold=True, color="17365D")

# Small footer note kept neutral and editable.
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("UNIVERSIAPOLIS Agadir • Génie Informatique • Année universitaire 2025–2026")
set_font(fr, size=8, color="7F8C8D")

# Document metadata.
doc.core_properties.title = "Demande de stage – El Houssaimi Sofiane"
doc.core_properties.subject = "Formalisation d’un stage en Génie Informatique"
doc.core_properties.author = "El Houssaimi Sofiane"
doc.core_properties.keywords = "stage, génie informatique, MikroTik, firewall, VPN"

doc.save(OUTPUT)
print(OUTPUT)

